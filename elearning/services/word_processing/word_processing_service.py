"""
Word Document Processing Service für DSP E-Learning Platform

Service für die Verarbeitung von Word-Dokumenten:
- Herunterladen von Word-Dokumenten aus Cloud Storage
- Extraktion von Text-Inhalt
- Verarbeitung durch WordExtraction Service
- JSON-Strukturierung für Datenbank-Speicherung

Author: DSP Development Team
Version: 1.0.0
"""
from __future__ import annotations
import io
import logging
from typing import Dict, Any, Optional, List
from dataclasses import dataclass

# Word Document Processing
try:
    from docx import Document
    WORD_PROCESSING_AVAILABLE = True
except ImportError:
    WORD_PROCESSING_AVAILABLE = False
    logging.warning("python-docx nicht verfügbar - Word Processing Service deaktiviert")

# Import des WordExtraction Services
from .word_extraction import WordExtraction

logger = logging.getLogger(__name__)


@dataclass
class ProcessedArticle:
    """Repräsentiert einen verarbeiteten Artikel."""
    title: str
    url: str
    json_content: Dict[str, Any]
    word_content: str
    file_name: str


class WordProcessingService:
    """
    Service für die Verarbeitung von Word-Dokumenten.
    
    Lädt Word-Dokumente herunter, extrahiert den Text-Inhalt
    und verarbeitet ihn durch den WordExtraction Service.
    """
    
    def __init__(self):
        self.word_extractor = WordExtraction()
        
        if not WORD_PROCESSING_AVAILABLE:
            logger.error("Word Processing nicht verfügbar - python-docx fehlt")
    
    def process_word_document(self, file_content: bytes, file_name: str, cloud_url: str) -> Optional[ProcessedArticle]:
        """
        Verarbeitet ein Word-Dokument und extrahiert strukturierten Inhalt.
        
        Args:
            file_content: Dateiinhalt als Bytes
            file_name: Name der Datei
            cloud_url: Cloud-URL der Datei
            
        Returns:
            ProcessedArticle mit extrahiertem Inhalt oder None
        """
        if not WORD_PROCESSING_AVAILABLE:
            logger.error("Word Processing nicht verfügbar")
            return None
        
        try:
            # Word-Dokument laden
            doc = Document(io.BytesIO(file_content))
            
            # Text-Inhalt extrahieren (inkl. Tabellen als Tabelle$-Blöcke in korrekter Reihenfolge)
            word_content = self._extract_text_from_docx(doc)
            
            # Titel aus Dateiname extrahieren
            title = self._extract_title_from_filename(file_name)
            
            # Durch WordExtraction verarbeiten
            print("[DEBUG] WordProcessingService: Übergabe an WordExtraction (mit Tabellenmarkern)")
            json_content = self.word_extractor.extract_content_to_json(word_content)
            
            # Tag-Analyse durchführen
            tag_analysis = self.word_extractor.analyze_tags_in_text(word_content)
            
            # JSON-Content mit Analyse erweitern
            json_content['tag_analysis'] = tag_analysis
            json_content['file_name'] = file_name
            json_content['cloud_url'] = cloud_url
            
            processed_article = ProcessedArticle(
                title=title,
                url=cloud_url,
                json_content=json_content,
                word_content=word_content,
                file_name=file_name
            )
            
            logger.info(f"Word-Dokument {file_name} erfolgreich verarbeitet")
            return processed_article
            
        except Exception as e:
            logger.error(f"Fehler beim Verarbeiten von {file_name}: {e}")
            return None
    
    def _extract_text_from_docx(self, doc: Document) -> str:
        """
        Extrahiert Text-Inhalt aus einem Word-Dokument in Original-Reihenfolge der Blöcke.
        Paragraphen werden als Zeilen ausgegeben, Tabellen werden als Tabelle$-Blöcke
        (mit Header + Zeilen) in den Text eingebettet, damit die Reihenfolge in der
        nachfolgenden Tag-Extraktion erhalten bleibt.
        
        Args:
            doc: Das geöffnete Word-Dokument
            
        Returns:
            Extrahierter Text als String
        """
        def _clean(s: str) -> str:
            return (s or "").replace("\u00ad", "").replace("\r", " ").replace("\n", " ").strip()

        # Helper: iterate body items in order (paragraphs and tables)
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        body = doc.element.body
        lines: list[str] = []

        for child in body.iterchildren():
            if isinstance(child, CT_P):
                p = Paragraph(child, doc)
                text = _clean(p.text)
                if text:
                    lines.append(text)
            elif isinstance(child, CT_Tbl):
                tbl = Table(child, doc)
                # Serialize table as Tabelle$ block
                # Build headers from first row (if available)
                headers: list[str] = []
                data_rows: list[list[str]] = []
                try:
                    if tbl.rows and tbl.columns:
                        headers = [_clean(c.text) for c in tbl.rows[0].cells]
                        for r in tbl.rows[1:]:
                            data_rows.append([_clean(c.text) for c in r.cells])
                except Exception:
                    # best-effort: skip malformed table
                    headers = []
                    data_rows = []

                if headers:
                    print(f"[DEBUG] _extract_text_from_docx: Tabelle gefunden mit {len(headers)} Spalten und {len(data_rows)} Zeilen")
                    lines.append("Tabelle$")
                    # join cells with pipe for robust splitting later
                    lines.append(" | ".join(headers))
                    for row in data_rows:
                        lines.append(" | ".join(row))
                    lines.append("Tabelle$")
                else:
                    # If table has no clear header, skip embedding
                    print("[DEBUG] _extract_text_from_docx: Tabelle übersprungen (keine Header erkannt)")
                    continue
            else:
                # Unhandled element types are ignored
                continue

        return "\n".join(lines)
    
    def _extract_title_from_filename(self, file_name: str) -> str:
        """
        Extrahiert einen Titel aus dem Dateinamen.
        
        Args:
            file_name: Name der Datei (z.B. "1.1 Installation und erste Schritte.docx")
            
        Returns:
            Extrahierter Titel
        """
        # Entferne .docx Erweiterung
        name_without_ext = file_name.replace('.docx', '')
        
        # Entferne Nummerierung am Anfang (z.B. "1.1 ")
        import re
        title = re.sub(r'^\d+\.\d+\s+', '', name_without_ext)
        
        return title if title else name_without_ext

    # Hinweis: separate Tabellen-Nachverarbeitung entfernt – Tabellen werden jetzt
    # in _extract_text_from_docx inline als Tabelle$-Blöcke codiert und in
    # WordExtraction an der korrekten Stelle ins JSON umgewandelt.
    
    def process_multiple_documents(self, documents: List[tuple]) -> List[ProcessedArticle]:
        """
        Verarbeitet mehrere Word-Dokumente.
        
        Args:
            documents: Liste von Tuples (file_content, file_name, cloud_url)
            
        Returns:
            Liste der verarbeiteten Artikel
        """
        processed_articles = []
        
        for file_content, file_name, cloud_url in documents:
            processed = self.process_word_document(file_content, file_name, cloud_url)
            if processed:
                processed_articles.append(processed)
        
        logger.info(f"{len(processed_articles)} von {len(documents)} Dokumenten erfolgreich verarbeitet")
        return processed_articles
    
    def validate_word_document(self, file_content: bytes, file_name: str) -> bool:
        """
        Validiert ein Word-Dokument.
        
        Args:
            file_content: Dateiinhalt als Bytes
            file_name: Name der Datei
            
        Returns:
            True wenn gültig, False sonst
        """
        if not WORD_PROCESSING_AVAILABLE:
            return False
        
        try:
            # Versuche Dokument zu öffnen
            doc = Document(io.BytesIO(file_content))
            
            # Prüfe ob Dokument Inhalt hat
            has_content = False
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    has_content = True
                    break
            
            if not has_content:
                logger.warning(f"Word-Dokument {file_name} hat keinen Inhalt")
                return False
            
            logger.info(f"Word-Dokument {file_name} ist gültig")
            return True
            
        except Exception as e:
            logger.error(f"Word-Dokument {file_name} ist ungültig: {e}")
            return False
    
    def get_document_statistics(self, file_content: bytes, file_name: str) -> Dict[str, Any]:
        """
        Erstellt Statistiken für ein Word-Dokument.
        
        Args:
            file_content: Dateiinhalt als Bytes
            file_name: Name der Datei
            
        Returns:
            Dictionary mit Statistiken
        """
        if not WORD_PROCESSING_AVAILABLE:
            return {}
        
        try:
            doc = Document(io.BytesIO(file_content))
            word_content = self._extract_text_from_docx(doc)
            
            # Tag-Analyse
            tag_analysis = self.word_extractor.analyze_tags_in_text(word_content)
            
            # Zusätzliche Statistiken
            stats = {
                'file_name': file_name,
                'file_size_bytes': len(file_content),
                'paragraphs_count': len(doc.paragraphs),
                'text_length': len(word_content),
                'word_count': len(word_content.split()),
                'tag_analysis': tag_analysis
            }
            
            return stats
            
        except Exception as e:
            logger.error(f"Fehler beim Erstellen der Statistiken für {file_name}: {e}")
            return {} 