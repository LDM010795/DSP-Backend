"""
Test-Script für detaillierte Tag-Analyse

Zeigt welche Tags gefunden und verarbeitet wurden.
"""

import sys
import os
import docx
import json

# Pfad zum elearning Modul hinzufügen
sys.path.append(os.path.join(os.path.dirname(__file__), 'elearning'))

from services.word_extraction import WordExtraction


def load_real_word_document(file_path: str) -> str:
    """
    Lädt echte Word-Datei und konvertiert zu Text.

    Args:
        file_path: Pfad zur .docx Datei

    Returns:
        str: Extrahierter Text
    """
    try:
        print(f"📄 Lade echte Word-Datei: {file_path}")
        print(f"📁 Datei existiert: {os.path.exists(file_path)}")

        doc = docx.Document(file_path)
        print(f"✅ Word-Dokument geladen, {len(doc.paragraphs)} Absätze gefunden")

        text = ""
        
        for i, paragraph in enumerate(doc.paragraphs):
            paragraph_text = paragraph.text.strip()
            if paragraph_text:
                # Jeder Word-Absatz wird als separate Zeile
                text += paragraph_text + "\n"
                print(f"   Absatz {i+1}: {paragraph_text[:80]}...")
            else:
                # Leere Absätze werden als leere Zeile behandelt
                text += "\n"

        print(f"✅ Text erfolgreich extrahiert ({len(text)} Zeichen)")
        return text

    except Exception as e:
        print(f"❌ Fehler beim Laden der Word-Datei: {e}")
        import traceback
        traceback.print_exc()
        return ""


def print_tag_analysis(analysis: dict):
    """
    Zeigt detaillierte Tag-Analyse an.

    Args:
        analysis: Tag-Analyse Dictionary
    """
    print("\n" + "="*80)
    print("📊 DETAILLIERTE TAG-ANALYSE")
    print("="*80)

    # Zusammenfassung
    summary = analysis['summary']
    print(f"📈 ZUSAMMENFASSUNG:")
    print(f"   • Gesamte Zeilen: {summary['total_lines']}")
    print(f"   • Verschiedene Tag-Arten gefunden: {summary['total_found_tags']}")
    print(f"   • Gesamte Tag-Vorkommen: {summary['total_found_occurrences']}")
    print(f"   • Verschiedene Tag-Arten verarbeitet: {summary['total_processed_tags']}")
    print(f"   • Gesamte verarbeitete Vorkommen: {summary['total_processed_occurrences']}")
    print(f"   • Verschiedene unbekannte Tag-Arten: {summary['total_unknown_tags']}")
    print(f"   • Gesamte unbekannte Vorkommen: {summary['total_unknown_occurrences']}")
    print(f"   • Ungenutzte Tags: {len(summary['unused_tags'])}")
    print(f"   • Unverarbeitete Tags: {len(summary['unprocessed_tags'])}")

    # Gefundene Tags mit Zählung
    print(f"\n🔍 GEFUNDENE TAGS (mit Anzahl):")
    for tag, count in analysis['found_tags_count'].items():
        print(f"   ✅ {tag}: {count}x")

    # Verarbeitete Tags mit Zählung
    print(f"\n✅ VERARBEITETE TAGS (mit Anzahl):")
    for tag, count in analysis['processed_tags_count'].items():
        print(f"   ✅ {tag}: {count}x")

    # Unbekannte Tags mit Zählung
    if analysis['unknown_tags_count']:
        print(f"\n❓ UNBEKANNTE TAGS (mögliche Tippfehler):")
        for tag, count in analysis['unknown_tags_count'].items():
            print(f"   ❓ {tag}: {count}x")

    # Ungenutzte Tags
    if analysis['unused_tags']:
        print(f"\n❌ UNGENUTZTE TAGS:")
        for tag in analysis['unused_tags']:
            print(f"   ❌ {tag}")

    # Unverarbeitete Tags
    if analysis['unprocessed_tags']:
        print(f"\n⚠️  UNVERARBEITETE TAGS (ohne Inhalt):")
        for tag in analysis['unprocessed_tags']:
            print(f"   ⚠️  {tag}")

    # Alle verfügbaren Tags
    print(f"\n📋 ALLE VERFÜGBAREN TAGS:")
    for tag in analysis['all_available_tags']:
        status = "✅" if tag in analysis['found_tags'] else "❌"
        count = analysis['found_tags_count'].get(tag, 0)
        if count > 0:
            print(f"   {status} {tag} ({count}x)")
        else:
            print(f"   {status} {tag}")


def save_json_to_file(data: dict, filename: str):
    """
    Speichert JSON in Datei.
    
    Args:
        data: JSON-Daten
        filename: Dateiname
    """
    try:
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        print(f"💾 JSON gespeichert in: {filename}")
    except Exception as e:
        print(f"❌ Fehler beim Speichern: {e}")


def main():
    """
    Hauptfunktion für die Tag-Analyse.
    """
    print("🧪 TEST: Detaillierte Tag-Analyse")
    print("="*60)
    
    # Word-Datei Pfad (im Root-Verzeichnis)
    word_file_path = "../1.1 Installation und erste Schritte.docx"
    
    # Prüfe ob Datei existiert
    if not os.path.exists(word_file_path):
        print(f"❌ Word-Datei nicht gefunden: {word_file_path}")
        print("💡 Stelle sicher, dass die Datei im Root-Verzeichnis liegt")
        return
    
    # 1. Word-Datei laden
    text = load_real_word_document(word_file_path)
    
    if not text:
        print("❌ Konnte Word-Datei nicht laden")
        return
    
    # 2. WordExtraction Service initialisieren
    print("\n🔄 Initialisiere WordExtraction Service...")
    extractor = WordExtraction()
    
    # 3. JSON-Extraktion durchführen
    print("🔍 Führe JSON-Extraktion durch...")
    json_result = extractor.extract_content_to_json(text)
    
    # 4. Tag-Analyse durchführen
    print("🔍 Führe Tag-Analyse durch...")
    analysis = extractor.analyze_tags_in_text(text)
    
    # 5. Analyse anzeigen
    print_tag_analysis(analysis)
    
    # 6. Beide JSON-Dateien speichern
    print("\n" + "="*60)
    print("💾 SPEICHERE JSON-DATEIEN")
    print("="*60)
    
    # JSON-Extraktion speichern
    save_json_to_file(json_result, "extracted_content.json")
    
    # Tag-Analyse speichern
    save_json_to_file(analysis, "tag_analysis.json")
    
    # 7. Zusammenfassung
    print("="*60)
    print("✅ TAG-ANALYSE ABGESCHLOSSEN")
    print("="*60)
    print(f"📊 JSON Content-Elemente: {len(json_result['content'])}")
    print(f"📊 Verschiedene Tag-Arten: {analysis['summary']['total_found_tags']}")
    print(f"📊 Gesamte Tag-Vorkommen: {analysis['summary']['total_found_occurrences']}")
    print(f"📊 Verarbeitete Tag-Arten: {analysis['summary']['total_processed_tags']}")
    print(f"📊 Verarbeitete Vorkommen: {analysis['summary']['total_processed_occurrences']}")
    print("📁 Dateien erstellt:")
    print("   - extracted_content.json (JSON Output)")
    print("   - tag_analysis.json (Detaillierte Analyse)")
    
    # 8. Empfehlungen
    print("\n💡 EMPFEHLUNGEN:")
    if len(analysis['summary']['unused_tags']) > 0:
        print("   • Ungenutzte Tags gefunden - prüfe ob alle Tags benötigt werden")
    if len(analysis['summary']['unprocessed_tags']) > 0:
        print("   • Unverarbeitete Tags gefunden - prüfe Tag-Struktur")
    
    # 9. Zeige erste paar JSON-Elemente
    if json_result['content']:
        print("\n🔍 ERSTE JSON-ELEMENTE:")
        for i, element in enumerate(json_result['content'][:3]):
            element_type = element.get('type', 'unknown')
            element_text = element.get('text', str(element))[:50] if 'text' in element else str(element)[:50]
            print(f"  {i+1}. {element_type}: {element_text}...")


if __name__ == "__main__":
    main() 